"""
Document Generator - Professional PDF & DOCX exports with industry-standard formatting.
Cinema-quality document output: Courier 12pt for screenplays, clean headers for other docs.
"""

import io
from datetime import datetime


BRAND = "Coffee with Cinema"


class DocumentGenerator:
    """Generates professionally formatted downloadable documents"""

    # -----------------------------------------------------------------------
    # PDF
    # -----------------------------------------------------------------------

    def generate_pdf(self, content: str, content_type: str) -> io.BytesIO:
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.lib import colors
            from reportlab.platypus import (
                SimpleDocTemplate, Paragraph, Spacer, HRFlowable,
                PageBreak, KeepTogether
            )
            from reportlab.platypus import Table, TableStyle

            buffer = io.BytesIO()
            PAGE_W, PAGE_H = letter

            # Margins — screenplay standard: 1.5" left, 1" rest
            left   = 1.5 * inch if content_type == 'screenplay' else 1.0 * inch
            right  = 1.0 * inch
            top    = 1.0 * inch
            bottom = 1.0 * inch

            doc = SimpleDocTemplate(
                buffer,
                pagesize=letter,
                leftMargin=left,
                rightMargin=right,
                topMargin=top,
                bottomMargin=bottom,
                title=content_type.replace('_', ' ').title(),
                author=BRAND,
            )

            styles = getSampleStyleSheet()

            # ---- Custom styles ------------------------------------------------
            COURIER = 'Courier'
            HELVETICA = 'Helvetica'
            HELVETICA_BOLD = 'Helvetica-Bold'

            title_style = ParagraphStyle(
                'CWCTitle',
                fontName=HELVETICA_BOLD,
                fontSize=20,
                textColor=colors.HexColor('#1a1a2e'),
                spaceAfter=6,
                leading=24,
            )
            subtitle_style = ParagraphStyle(
                'CWCSubtitle',
                fontName=HELVETICA,
                fontSize=10,
                textColor=colors.HexColor('#555555'),
                spaceAfter=20,
            )
            # Screenplay: heading (INT./EXT.)
            slug_style = ParagraphStyle(
                'SlugLine',
                fontName=COURIER,
                fontSize=12,
                textColor=colors.black,
                spaceBefore=18,
                spaceAfter=6,
                leading=14.4,
            )
            # Screenplay: action lines
            action_style = ParagraphStyle(
                'Action',
                fontName=COURIER,
                fontSize=12,
                textColor=colors.black,
                spaceAfter=12,
                leading=14.4,
            )
            # Screenplay: character name (centered, ~3.7" from left)
            char_style = ParagraphStyle(
                'CharName',
                fontName=COURIER,
                fontSize=12,
                textColor=colors.black,
                leftIndent=2.2 * inch,
                spaceBefore=12,
                leading=14.4,
            )
            # Screenplay: dialogue (~2.5" from left, max ~3.5" wide)
            dialogue_style = ParagraphStyle(
                'Dialogue',
                fontName=COURIER,
                fontSize=12,
                textColor=colors.black,
                leftIndent=1.3 * inch,
                rightIndent=1.3 * inch,
                spaceAfter=6,
                leading=14.4,
            )
            # Screenplay: parenthetical
            paren_style = ParagraphStyle(
                'Paren',
                fontName=COURIER,
                fontSize=12,
                textColor=colors.HexColor('#333333'),
                leftIndent=1.7 * inch,
                spaceAfter=2,
                leading=14.4,
            )
            # Transition (right-aligned)
            trans_style = ParagraphStyle(
                'Trans',
                fontName=COURIER,
                fontSize=12,
                textColor=colors.black,
                alignment=2,  # RIGHT
                spaceBefore=12,
                spaceAfter=6,
            )
            # General section heading (non-screenplay docs)
            section_style = ParagraphStyle(
                'Section',
                fontName=HELVETICA_BOLD,
                fontSize=13,
                textColor=colors.HexColor('#1a1a2e'),
                spaceBefore=20,
                spaceAfter=6,
                borderPad=4,
            )
            divider_style = ParagraphStyle(
                'Divider',
                fontName=HELVETICA_BOLD,
                fontSize=11,
                textColor=colors.HexColor('#c4a484'),
                spaceBefore=14,
                spaceAfter=4,
            )
            body_style = ParagraphStyle(
                'CWCBody',
                fontName=COURIER,
                fontSize=10,
                textColor=colors.HexColor('#111111'),
                spaceAfter=8,
                leading=15,
            )

            # ---- Build story -----------------------------------------------
            story = []

            # Cover block
            cfont_display = content_type.replace('_', ' ').title()
            story.append(Spacer(1, 0.3 * inch))
            story.append(Paragraph(cfont_display, title_style))
            story.append(Paragraph(
                f"{BRAND}  ·  Generated {datetime.now().strftime('%B %d, %Y')}",
                subtitle_style
            ))
            story.append(HRFlowable(
                width="100%", thickness=1,
                color=colors.HexColor('#c4a484'),
                spaceAfter=20
            ))

            # ---- Screenplay parse ------------------------------------------
            if content_type == 'screenplay':
                in_dialogue = False
                lines = content.split('\n')
                i = 0
                while i < len(lines):
                    raw = lines[i]
                    stripped = raw.strip()

                    if not stripped:
                        i += 1
                        continue

                    upper = stripped.upper()

                    # Transitions
                    if upper in ('FADE IN:', 'FADE OUT.', 'CUT TO:', 'DISSOLVE TO:', 'SMASH CUT TO:', 'MATCH CUT TO:'):
                        story.append(Paragraph(stripped, trans_style))
                        in_dialogue = False

                    # Scene headings
                    elif (stripped.startswith('INT.') or stripped.startswith('EXT.')
                          or stripped.startswith('INT ') or stripped.startswith('EXT ')):
                        story.append(Spacer(1, 0.1 * inch))
                        story.append(Paragraph(stripped.upper(), slug_style))
                        in_dialogue = False

                    # Character name (ALL CAPS, 1-4 words, next line probably dialogue)
                    elif (stripped == stripped.upper()
                          and len(stripped.split()) <= 5
                          and len(stripped) < 40
                          and stripped.isalpha() or (stripped.replace(' ', '').replace("'", '').replace('-', '').isalpha())):
                        story.append(Paragraph(stripped, char_style))
                        in_dialogue = True

                    # Parenthetical
                    elif stripped.startswith('(') and stripped.endswith(')'):
                        story.append(Paragraph(stripped, paren_style))

                    # Dialogue (after character name)
                    elif in_dialogue and not stripped.startswith('INT.') and not stripped.startswith('EXT.'):
                        story.append(Paragraph(stripped, dialogue_style))

                    # Action / description
                    else:
                        story.append(Paragraph(stripped, action_style))
                        in_dialogue = False

                    i += 1

            # ---- Non-screenplay documents ---------------------------------
            else:
                lines = content.split('\n')
                for line in lines:
                    stripped = line.strip()
                    if not stripped:
                        story.append(Spacer(1, 6))
                        continue

                    # Detect section headers (═══ or === or ALL CAPS header)
                    if stripped.startswith('═') or stripped.startswith('===') or stripped.startswith('---'):
                        story.append(HRFlowable(width="100%", thickness=0.5,
                                                 color=colors.HexColor('#c4a484'),
                                                 spaceBefore=8, spaceAfter=8))
                    elif (stripped == stripped.upper()
                          and len(stripped) > 4
                          and len(stripped) < 80
                          and not stripped.startswith('•')
                          and not stripped.startswith('-')):
                        story.append(Paragraph(stripped, divider_style))
                    elif stripped.startswith('•') or stripped.startswith('*') or stripped.startswith('-'):
                        clean = stripped.lstrip('•*- ').strip()
                        story.append(Paragraph(f"• {clean}", body_style))
                    else:
                        # Check if it's a subsection heading (ends with :)
                        if stripped.endswith(':') and len(stripped) < 60 and stripped[0].isupper():
                            story.append(Paragraph(stripped, section_style))
                        else:
                            story.append(Paragraph(stripped, body_style))

            doc.build(story)
            buffer.seek(0)
            return buffer

        except ImportError:
            buf = io.BytesIO(content.encode('utf-8'))
            buf.seek(0)
            return buf

    # -----------------------------------------------------------------------
    # DOCX
    # -----------------------------------------------------------------------

    def generate_docx(self, content: str, content_type: str) -> io.BytesIO:
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            import copy

            doc = Document()

            # Page margins
            for section in doc.sections:
                section.page_width  = Inches(8.5)
                section.page_height = Inches(11)
                if content_type == 'screenplay':
                    section.left_margin   = Inches(1.5)
                    section.right_margin  = Inches(1.0)
                else:
                    section.left_margin   = Inches(1.2)
                    section.right_margin  = Inches(1.0)
                section.top_margin    = Inches(1.0)
                section.bottom_margin = Inches(1.0)

            GOLD = RGBColor(0xC4, 0xA4, 0x84)
            DARK = RGBColor(0x1A, 0x1A, 0x2E)

            # --- Title block ---
            title_p = doc.add_paragraph()
            title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = title_p.add_run(content_type.replace('_', ' ').upper())
            run.bold = True
            run.font.size = Pt(22)
            run.font.color.rgb = DARK
            run.font.name = 'Helvetica' if content_type != 'screenplay' else 'Courier New'

            sub_p = doc.add_paragraph()
            sub_r = sub_p.add_run(
                f"{BRAND}  ·  {datetime.now().strftime('%B %d, %Y')}"
            )
            sub_r.font.size = Pt(9)
            sub_r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

            # Decorative rule
            hr = doc.add_paragraph()
            hr_run = hr.add_run('─' * 70)
            hr_run.font.color.rgb = GOLD
            hr_run.font.size = Pt(9)

            doc.add_paragraph()

            # --- Content ---
            if content_type == 'screenplay':
                in_dialogue = False
                for raw in content.split('\n'):
                    stripped = raw.strip()
                    if not stripped:
                        doc.add_paragraph()
                        continue

                    upper = stripped.upper()

                    if upper in ('FADE IN:', 'FADE OUT.', 'CUT TO:', 'DISSOLVE TO:'):
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        r = p.add_run(stripped)
                        r.font.name = 'Courier New'
                        r.font.size = Pt(12)
                        in_dialogue = False

                    elif stripped.startswith(('INT.', 'EXT.', 'INT ', 'EXT ')):
                        p = doc.add_paragraph()
                        r = p.add_run(stripped.upper())
                        r.bold = True
                        r.font.name = 'Courier New'
                        r.font.size = Pt(12)
                        pf = p.paragraph_format
                        pf.space_before = Pt(18)
                        in_dialogue = False

                    elif (stripped == stripped.upper()
                          and len(stripped.split()) <= 5
                          and len(stripped) < 40):
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        r = p.add_run(stripped)
                        r.font.name = 'Courier New'
                        r.font.size = Pt(12)
                        p.paragraph_format.space_before = Pt(12)
                        in_dialogue = True

                    elif stripped.startswith('(') and stripped.endswith(')'):
                        p = doc.add_paragraph()
                        p.paragraph_format.left_indent = Inches(1.5)
                        r = p.add_run(stripped)
                        r.font.name = 'Courier New'
                        r.font.size = Pt(12)

                    elif in_dialogue:
                        p = doc.add_paragraph()
                        p.paragraph_format.left_indent  = Inches(1.0)
                        p.paragraph_format.right_indent = Inches(1.0)
                        r = p.add_run(stripped)
                        r.font.name = 'Courier New'
                        r.font.size = Pt(12)

                    else:
                        p = doc.add_paragraph()
                        r = p.add_run(stripped)
                        r.font.name = 'Courier New'
                        r.font.size = Pt(12)
                        in_dialogue = False

            else:
                for raw in content.split('\n'):
                    stripped = raw.strip()
                    if not stripped:
                        doc.add_paragraph()
                        continue

                    if stripped.startswith('═') or stripped.startswith('===') or stripped.startswith('---'):
                        p = doc.add_paragraph()
                        r = p.add_run('─' * 60)
                        r.font.color.rgb = GOLD
                        r.font.size = Pt(8)

                    elif (stripped == stripped.upper()
                          and 4 < len(stripped) < 80
                          and not stripped.startswith('•')
                          and not stripped.startswith('-')):
                        p = doc.add_paragraph()
                        r = p.add_run(stripped)
                        r.bold = True
                        r.font.size = Pt(13)
                        r.font.color.rgb = GOLD
                        r.font.name = 'Helvetica'
                        p.paragraph_format.space_before = Pt(16)

                    elif stripped.startswith('•') or stripped.startswith('*') or stripped.startswith('-'):
                        clean = stripped.lstrip('•*- ').strip()
                        p = doc.add_paragraph(style='List Bullet')
                        r = p.add_run(clean)
                        r.font.size = Pt(11)

                    elif stripped.endswith(':') and len(stripped) < 60 and stripped[0].isupper():
                        p = doc.add_paragraph()
                        r = p.add_run(stripped)
                        r.bold = True
                        r.font.size = Pt(11)
                        r.font.color.rgb = DARK

                    else:
                        p = doc.add_paragraph()
                        r = p.add_run(stripped)
                        r.font.size = Pt(11)

            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer

        except ImportError:
            buf = io.BytesIO(content.encode('utf-8'))
            buf.seek(0)
            return buf
